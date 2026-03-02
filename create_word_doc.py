"""
Script to create a Word document user guide from the ROW Application PPT content.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Paths
OUTPUT_DOC = r"C:\Users\Ajitha.Rajkumar\chatbot\row\ROW_Application_User_Guide.docx"
IMAGES_DIR = r"C:\Users\Ajitha.Rajkumar\chatbot\row\extracted\images"


def img(filename):
    """Helper to get full image path."""
    return os.path.join(IMAGES_DIR, filename)


def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading


def add_body_text(doc, text, bold=False, italic=False):
    """Add a paragraph with consistent body styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    return p


def add_step(doc, step_num, text):
    """Add a numbered step."""
    p = doc.add_paragraph()
    run_num = p.add_run(f"Step {step_num}: ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.name = "Calibri"
    run_num.font.color.rgb = RGBColor(0, 51, 102)
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)
    run_text.font.name = "Calibri"
    return p


def add_screenshot(doc, image_path, caption="", width=Inches(5.8)):
    """Add a screenshot image with an optional caption."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = cap.add_run(caption)
            run_cap.font.size = Pt(9)
            run_cap.font.italic = True
            run_cap.font.color.rgb = RGBColor(100, 100, 100)


def add_note(doc, text):
    """Add a note/tip box."""
    p = doc.add_paragraph()
    run_label = p.add_run("Note: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(204, 102, 0)
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.italic = True
    return p


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table


# ============================================================
# CREATE THE DOCUMENT
# ============================================================
doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

title = doc.add_heading("ROW Application", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_heading("User Guide", level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph("")
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Infra Build – Jio Digital Life")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("")
desc2 = doc.add_paragraph()
desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = desc2.add_run("A step-by-step guide for ROW Executives to create, manage,\nmerge, and submit ROW Applications using the K2 and GIS systems.")
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled(doc, "Table of Contents", level=1)
toc_items = [
    "1. Overview",
    "2. Workflow Summary",
    "3. Accessing the ROW Application",
    "4. Searching and Selecting Span Details",
    "5. Creating a ROW Application",
    "6. SPAN Split and Update Authority",
    "7. Saving the Split Span",
    "8. Updating Authority for Split Spans",
    "9. Generating Application After Split",
    "10. Merging Applications",
    "11. Application Acknowledgement & Submission",
    "12. Capture Joint Survey Details",
    "13. Capture Application Processing Fee",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"

doc.add_page_break()

# ============================================================
# 1. OVERVIEW
# ============================================================
add_heading_styled(doc, "1. Overview", level=1)

add_body_text(doc,
    "The ROW (Right of Way) Application is a web-based module within the Infra Build platform "
    "(Jio Digital Life) that enables ROW Executives to create, manage, and submit Right of Way "
    "applications for fiber optic network deployments. The application integrates with K2 workflow "
    "engine and GIS (Geographic Information System) for geographic span management."
)

add_body_text(doc,
    "This document provides a comprehensive step-by-step guide covering the complete lifecycle of "
    "a ROW Application — from initial creation through span management, application merging, "
    "submission, joint survey, and processing fee capture."
)

add_body_text(doc, "Key Capabilities:", bold=True)
bullets = [
    "Create ROW Draft Applications for Intercity & Intracity networks",
    "View and select SPAN, CROSSING, and POLE details from GIS",
    "Split spans and update authority information",
    "Merge multiple draft applications into a single application",
    "Submit application acknowledgement with supporting documents",
    "Capture Joint Survey findings",
    "Capture and approve Application Processing Fee (multi-level approval)",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_page_break()

# ============================================================
# 2. WORKFLOW SUMMARY
# ============================================================
add_heading_styled(doc, "2. Workflow Summary", level=1)

add_body_text(doc,
    "The following table summarizes the key tasks in the ROW Application workflow:"
)

add_table(doc,
    ["Workflow", "Task Name", "System", "App Type", "Role", "Location"],
    [
        ["ROW Application", "Create ROW Application", "K2, GIS", "Web", "ROW Executive", "SHQ"],
        ["ROW Application", "Application Acknowledgement", "K2, GIS", "Web", "ROW Executive", "SHQ"],
        ["ROW Application", "Capture Joint Survey Details", "K2, GIS", "Web", "ROW Executive", "SHQ"],
        ["ROW Application", "Capture Application Processing Fee", "K2, GIS", "Web", "ROW Executive", "SHQ"],
        ["ROW Application", "Processing Fee – Approval L1", "K2, GIS", "Web", "SCO", "SHQ"],
        ["ROW Application", "Processing Fee – Approval L2", "K2, GIS", "Web", "State FC&A", "SHQ"],
    ]
)

doc.add_page_break()

# ============================================================
# 3. ACCESSING THE ROW APPLICATION
# ============================================================
add_heading_styled(doc, "3. Accessing the ROW Application", level=1)

add_body_text(doc,
    "To begin working with ROW Applications, you must first access the FTTX Home screen on the K2 platform."
)

add_step(doc, 1, 'Log in to the K2 platform and navigate to the FTTX Home screen.')
add_step(doc, 2, 'On the FTTX Home screen, locate the "PUBLIC ROW" tile on the left side of the screen.')
add_step(doc, 3, 'Click on "Manage Draft Application (Intercity & Intracity)" under the PUBLIC ROW section.')

add_screenshot(doc, img("slide3_image1.png"),
    "Figure 1: FTTX Home Screen – Click 'Manage Draft Application (Intercity & Intracity)' under PUBLIC ROW")

add_note(doc, "The PUBLIC ROW tile contains several options including Manage Authority, ROW Payment Doc Status Updation, "
    "ROW Manual Payment Doc Status Updation, Manage Application OneFiber, Manage Draft Application (FTTX/B2O), "
    "Manage Demand Note, and Extend Refundables.")

doc.add_page_break()

# ============================================================
# 4. SEARCHING AND SELECTING SPAN DETAILS
# ============================================================
add_heading_styled(doc, "4. Searching and Selecting Span Details", level=1)

add_body_text(doc,
    "After accessing the application, you will be directed to the 'Application Generation One Fiber' screen "
    "where you can search for span details."
)

add_step(doc, 1, 'Select "NON DIT Draft Application" as the application type.')
add_step(doc, 2, 'Fill in the search filters:')

filter_items = [
    "Phase – Select the applicable phase",
    "Jio State – e.g., Chhattisgarh",
    "Political State – e.g., Chhattisgarh",
    "MP (Maintenance Point) – e.g., INCGJGDP01 - Jagdalpur",
    "Authority Category – e.g., NHAI",
    "Authority – e.g., CG NHAI",
    "Network Category – Select as applicable",
    "Network Type – e.g., INTERCITY",
]
for item in filter_items:
    doc.add_paragraph(item, style="List Bullet")

add_step(doc, 3, 'Click the "Search" button to retrieve the link/route details.')
add_step(doc, 4, 'In the LINKS/SPANS section, select the Link/Route from the dropdown (e.g., CG_CG_CGMADIORUMTESPR002_BU).')
add_step(doc, 5, 'Click "Get Span" to retrieve span details. You can also click "Get Pole" or "Get Crossing" to get pole and crossing details.')
add_step(doc, 6, 'SPAN, CROSSING, and POLE details eligible for ROW will be displayed in a table showing Feature ID, Length, UOM, Construction Type, and Construction Methodology.')

add_screenshot(doc, img("slide4_image2.png"),
    "Figure 2: Application Generation One Fiber – Search filters and SPAN details")

doc.add_page_break()

# ============================================================
# 5. CREATING A ROW APPLICATION
# ============================================================
add_heading_styled(doc, "5. Creating a ROW Application", level=1)

add_body_text(doc,
    "Once you have retrieved the span details, you can select the desired spans and generate a ROW application."
)

add_step(doc, 1, 'Select the desired SPAN segment(s) by checking the checkbox next to the Feature ID.')
add_step(doc, 2, 'The "Selected Span Length (Mtrs)" field will automatically update to reflect the total length of selected spans.')
add_step(doc, 3, 'Click the "Generate Application" button at the bottom of the page.')
add_step(doc, 4, 'A success message will appear confirming the application has been created, displaying the Application Reference Number (e.g., NLD-CG-CG-2022034154204843).')
add_step(doc, 5, 'Click "OK" to dismiss the success dialog.')

add_screenshot(doc, img("slide5_image3.png"),
    "Figure 3: Application Created Successfully – Application Reference Number displayed")

doc.add_page_break()

# ============================================================
# 6. SPAN SPLIT AND UPDATE AUTHORITY
# ============================================================
add_heading_styled(doc, "6. SPAN Split and Update Authority", level=1)

add_body_text(doc,
    "Before application creation, users can split spans and update the authority information. "
    "This is done through the Update Authority page integrated with GIS mapping."
)

add_heading_styled(doc, "6.1 Viewing Authority Details", level=2)

add_body_text(doc,
    "You can view and manage all registered authorities from the Authority Details page."
)

add_step(doc, 1, 'Navigate to the Authority Details page in Infra Build.')
add_step(doc, 2, 'The page displays a list of all authorities with their Category, Name, Type, Number of Circles, MPs, JPs, and Address.')
add_step(doc, 3, 'Use the toolbar options: Add Authority, Edit Authority, Manage Authority Cost, Manage Authority Vendor Contract Type Mapping, or Update ROW Authority (GIS).')

add_screenshot(doc, img("slide7_image5.png"),
    "Figure 4: Authority Details – List of all registered authorities")

add_heading_styled(doc, "6.2 Adding/Updating Authority", level=2)

add_body_text(doc,
    "If the authority for a span is not registered, you can add a new authority."
)

add_step(doc, 1, 'Select the Authority Type: "Existing" or "New".')
add_step(doc, 2, 'For a new (unregistered) authority, fill in:')

auth_items = [
    "Authority Category – e.g., PWD",
    "Authority Name – e.g., CG PWD",
    "Authority Address – e.g., CG PWD MAIN LANE CG",
]
for item in auth_items:
    doc.add_paragraph(item, style="List Bullet")

add_step(doc, 3, 'Configure the JIO State Mapping (Circle and MP Name).')
add_step(doc, 4, 'Click "Save Authority" to save the new authority.')
add_step(doc, 5, 'A confirmation dialog will appear: "Are you sure you want to add/update authority with below mapping?" Click "OK".')
add_step(doc, 6, 'A success message will confirm: "Authority Saved Successfully."')

add_screenshot(doc, img("slide8_image6.png"),
    "Figure 5: Add New Authority – Authority Type, Category, Name, and State Mapping")

add_screenshot(doc, img("slide9_image7.png"),
    "Figure 6: Authority Saved Successfully")

doc.add_page_break()

# ============================================================
# 6.3 USING GIS MAP FOR SPAN MANAGEMENT
# ============================================================
add_heading_styled(doc, "6.3 Using GIS Map for Span Management", level=2)

add_body_text(doc,
    "The Update Authority from GIS Page provides an interactive map view where you can visualize "
    "fiber spans and perform split/merge operations."
)

add_step(doc, 1, 'On the GIS page, fill in the filter fields:')
gis_items = [
    "Jio State and State – e.g., Chhattisgarh",
    "Maintenance Point – e.g., Jagdalpur",
    "Network Type – e.g., INTERCITY",
    "Link Id and Span Id – Select from dropdowns",
    "Feature Type – e.g., Span",
]
for item in gis_items:
    doc.add_paragraph(item, style="List Bullet")

add_step(doc, 2, 'Click "Submit" to load the span on the GIS map.')
add_step(doc, 3, 'The map will display the fiber route with color-coded spans. The span details (RWS Id, Length, Authority Category, Authority Name) are shown in the table below the map.')
add_step(doc, 4, 'Use the action buttons to manage spans:')

action_items = [
    "Start Split – Begin splitting a selected span into multiple segments",
    "Merge Span – Merge two or more span segments",
    "Stop Split – Stop the current split operation",
    "Reset Split – Reset the split to the original state",
    "ReSplit – Re-split a previously split span",
]
for item in action_items:
    doc.add_paragraph(item, style="List Bullet")

add_screenshot(doc, img("slide10_image8.png"),
    "Figure 7: GIS Map View – Span segments with color-coded routes and action buttons")

doc.add_page_break()

# ============================================================
# 7. SAVING THE SPLIT SPAN
# ============================================================
add_heading_styled(doc, "7. Saving the Split Span", level=1)

add_body_text(doc,
    "After splitting a span on the GIS map, you need to save the split data."
)

add_step(doc, 1, 'After performing the split operation on the map, the split spans will appear in the table with updated RWS IDs and lengths.')
add_step(doc, 2, 'Verify the split span details (RWS Id, Length in Mtrs, Authority Category, Authority Name).')
add_step(doc, 3, 'Click the "Save Split" button to save the split data.')
add_step(doc, 4, 'A success message will confirm: "Split span data saved successfully." Click "OK".')

add_screenshot(doc, img("slide12_image10.png"),
    "Figure 8: Split Span on Map – RWS IDs and lengths of split segments")

add_screenshot(doc, img("slide13_image11.png"),
    "Figure 9: Split Span Data Saved Successfully")

doc.add_page_break()

# ============================================================
# 8. UPDATING AUTHORITY FOR SPLIT SPANS
# ============================================================
add_heading_styled(doc, "8. Updating Authority for Split Spans", level=1)

add_body_text(doc,
    "After splitting a span, you may need to update the authority for individual split segments."
)

add_step(doc, 1, 'On the GIS map page, locate the split span segments in the table.')
add_step(doc, 2, 'For each span segment, select the appropriate Authority Category (e.g., NHAI, PWD) from the dropdown.')
add_step(doc, 3, 'Select the corresponding Authority Name (e.g., CG NHAI, CG PWD).')
add_step(doc, 4, 'Click the "Update" button next to the span segment to save the authority change.')
add_step(doc, 5, 'A success message will confirm: "Span data updated successfully." Click "OK".')

add_screenshot(doc, img("slide14_image12.png"),
    "Figure 10: Update Authority – Select Authority Category and Name for each split span")

add_screenshot(doc, img("slide6_image4.png"),
    "Figure 11: Span Data Updated Successfully")

doc.add_page_break()

# ============================================================
# 9. GENERATING APPLICATION AFTER SPLIT
# ============================================================
add_heading_styled(doc, "9. Generating Application After Split", level=1)

add_body_text(doc,
    "After splitting spans and updating authorities, you can return to the Application Generation "
    "screen to create a new application with the updated span details."
)

add_step(doc, 1, 'Navigate back to the "Application Generation One Fiber" screen.')
add_step(doc, 2, 'Search for spans using the same filters. You will now see the split span segments (e.g., INCGJGDP01_RWS0229 – 866.89 M, INCGJGDP01_RWS0258 – 418.81 M, INCGJGDP01_RWS0259 – 445.06 M).')
add_step(doc, 3, 'Select the desired span segments by checking the checkboxes.')
add_step(doc, 4, 'Click "Generate Application" to create the application.')
add_step(doc, 5, 'A success message will display the new Application Reference Number (e.g., NLD-CG-CG-2022034155725720). Click "OK".')

add_screenshot(doc, img("slide15_image13.png"),
    "Figure 12: Application Generation – Select split spans and generate application")

add_screenshot(doc, img("slide16_image14.png"),
    "Figure 13: Second Application Created Successfully after Span Split")

doc.add_page_break()

# ============================================================
# 10. MERGING APPLICATIONS
# ============================================================
add_heading_styled(doc, "10. Merging Applications", level=1)

add_body_text(doc,
    "Users can merge multiple draft applications into a single application. "
    "This is useful when multiple span segments under the same authority need "
    "to be submitted as one combined application."
)

add_step(doc, 1, 'Navigate to the "Manage Application OneFiber" option on the FTTX Home screen under PUBLIC ROW.')

add_screenshot(doc, img("slide18_image16.png"),
    "Figure 14: FTTX Home – Click 'Manage Application OneFiber'")

add_step(doc, 2, 'On the Merge Application screen, select the Network Type (e.g., INTERCITY).')
add_step(doc, 3, 'Fill in the Search Application filters (JIO State, Political State, MP, Authority Category, Authority).')
add_step(doc, 4, 'Click "Search" to retrieve draft applications.')
add_step(doc, 5, 'In the DRAFT APPLICATION section, select the application(s) to merge by checking the checkboxes.')

add_screenshot(doc, img("slide19_image17.png"),
    "Figure 15: Merge Application – Search and select draft applications")

add_step(doc, 6, 'The SELECTED DRAFT APPLICATIONS section will display the chosen applications with details including Site Name, Application Ref No, App Length, Pole Count, Crossing Count, Network Type, and Authority.')
add_step(doc, 7, 'Review the following details:')

merge_items = [
    "Selected App Length (Mtrs) – Total combined length",
    "Row Processing Fee Applicable – Select Yes/No",
    "App Pole (Nos) – Total pole count",
    "App Crossing (Nos) – Total crossing count",
    "Remarks – Add any remarks if needed",
]
for item in merge_items:
    doc.add_paragraph(item, style="List Bullet")

add_step(doc, 8, 'Click "Generate Merged Application" to create the merged application.')
add_step(doc, 9, 'A success message will confirm: "Merge Application (CG22034181336766) Generated successfully." Click "OK".')

add_screenshot(doc, img("slide20_image18.png"),
    "Figure 16: Merge Application – Review details and generate merged application")

add_screenshot(doc, img("slide21_image19.png"),
    "Figure 17: Merge Application Generated Successfully")

doc.add_page_break()

# ============================================================
# 11. APPLICATION ACKNOWLEDGEMENT & SUBMISSION
# ============================================================
add_heading_styled(doc, "11. Application Acknowledgement & Submission", level=1)

add_body_text(doc,
    "After the application is generated, a task will be assigned to the user for application submission. "
    "The user must complete the acknowledgement process."
)

add_step(doc, 1, 'Navigate to the Task List in Infra Build (FTTXTaskList).')
add_step(doc, 2, 'In the Task List, locate the task "Submit App Acknowledgement Details" under the ROW - Application process.')

add_screenshot(doc, img("slide22_image20.png"),
    "Figure 18: Task List – 'Submit App Acknowledgement Details' task assigned")

add_step(doc, 3, 'Click on the task to open the Submit Application Acknowledgement Details form.')
add_step(doc, 4, 'On the form, review the application details:')

ack_items = [
    "App. Ref. No. – Application reference number",
    "Acknowledgement ID – Will be assigned after submission",
    "Authority Name – e.g., CG NHAI",
    "Application Length (Mtrs) – e.g., 1285.70",
    "Application Pole (Nos) – Count of poles",
    "Application Crossing (Nos) – Count of crossings",
    "Submitted Date – Date of submission",
    "Submitted By – Name of the submitter",
]
for item in ack_items:
    doc.add_paragraph(item, style="List Bullet")

add_step(doc, 5, 'Upload the acknowledged copy of the submitted application (PDF format).')
add_step(doc, 6, 'Attach any Other Documents if needed.')
add_step(doc, 7, 'Enter Processing Fee (Rs.) if applicable.')
add_step(doc, 8, 'Select Payment Type and enter Payment Instrument No.')
add_step(doc, 9, 'Upload Processing Fee Scanned Copy if applicable.')
add_step(doc, 10, 'Click "Submit" to submit the application.')
add_step(doc, 11, 'A success message will confirm: "Application submitted." Click "OK".')

add_screenshot(doc, img("slide23_image21.png"),
    "Figure 19: Submit Application Acknowledgement – Application Submitted Successfully")

doc.add_page_break()

# ============================================================
# 12. CAPTURE JOINT SURVEY DETAILS
# ============================================================
add_heading_styled(doc, "12. Capture Joint Survey Details", level=1)

add_body_text(doc,
    "After the application is submitted, a task for capturing Joint Survey findings will be assigned."
)

add_step(doc, 1, 'Navigate to the Task List and locate the task "Submit Joint Survey Findings" under ROW - Application.')

add_screenshot(doc, img("slide24_image22.png"),
    "Figure 20: Task List – 'Submit Joint Survey Findings' task assigned")

add_step(doc, 2, 'Click the task to open the Joint Survey Findings form.')
add_step(doc, 3, 'On the form, review the application details:')

survey_items = [
    "App. Ref. No. – Application reference number (e.g., CG22034181336766)",
    "Authority Name – e.g., CG NHAI",
    "App. Length (Mtrs) – e.g., 1285.70",
    "App. Crossing (Nos) – Count of crossings",
    "App. Pole (Nos) – Count of poles",
]
for item in survey_items:
    doc.add_paragraph(item, style="List Bullet")

add_step(doc, 4, 'Enter the Date of Joint Survey.')
add_step(doc, 5, 'Upload the Survey Finding Document (e.g., Certificate of Completion PDF).')
add_step(doc, 6, 'Add Comments (e.g., "Joint survey done").')
add_step(doc, 7, 'Click "Submit" to submit the joint survey findings. You can also click "Not Applicable" if the survey is not required.')
add_step(doc, 8, 'A success message will confirm: "Joint Survey Findings Submitted." Click "OK".')

add_screenshot(doc, img("slide25_image23.png"),
    "Figure 21: Joint Survey Findings Form – Enter survey details and upload documents")

add_screenshot(doc, img("slide26_image24.png"),
    "Figure 22: Joint Survey Findings Submitted Successfully")

doc.add_page_break()

# ============================================================
# 13. CAPTURE APPLICATION PROCESSING FEE
# ============================================================
add_heading_styled(doc, "13. Capture Application Processing Fee", level=1)

add_body_text(doc,
    "The final step in the ROW Application workflow is capturing the Application Processing Fee. "
    "This task involves a multi-level approval process."
)

add_body_text(doc, "Approval Workflow:", bold=True)

add_table(doc,
    ["Workflow", "Task Name", "System", "App Type", "Role", "Location"],
    [
        ["ROW Application", "Capture Application Processing Fee", "K2, GIS", "Web", "ROW Executive", "SHQ"],
        ["ROW Application", "Processing Fee – Approval L1", "K2, GIS", "Web", "SCO", "SHQ"],
        ["ROW Application", "Processing Fee – Approval L2", "K2, GIS", "Web", "State FC&A", "SHQ"],
    ]
)

doc.add_paragraph("")

add_step(doc, 1, 'The ROW Executive captures the Application Processing Fee details.')
add_step(doc, 2, 'The fee is then sent for Level 1 (L1) Approval to the SCO (State Circle Office).')
add_step(doc, 3, 'After L1 approval, the fee goes for Level 2 (L2) Approval to the State FC&A (Finance, Costing & Accounts).')
add_step(doc, 4, 'Once both levels of approval are completed, the processing fee is finalized.')

add_note(doc,
    "The processing fee approval follows a hierarchical workflow. The ROW Executive initiates the fee capture, "
    "the SCO reviews and approves at L1, and the State FC&A provides final approval at L2."
)

doc.add_page_break()

# ============================================================
# APPENDIX
# ============================================================
add_heading_styled(doc, "Appendix", level=1)

add_heading_styled(doc, "A. Glossary", level=2)

glossary = [
    ("ROW", "Right of Way – Permission required to lay fiber optic cables along public roads/highways"),
    ("K2", "Workflow management platform used for task routing and approvals"),
    ("GIS", "Geographic Information System – Used for mapping and managing geographic span data"),
    ("FTTX", "Fiber to the X – Generic term for fiber optic network architecture"),
    ("SPAN", "A segment of fiber optic cable between two points"),
    ("RWS", "ROW Span – Unique identifier for span segments"),
    ("NHAI", "National Highways Authority of India"),
    ("PWD", "Public Works Department"),
    ("DIT", "Department of Information Technology"),
    ("SCO", "State Circle Office"),
    ("FC&A", "Finance, Costing & Accounts"),
    ("MP", "Maintenance Point"),
    ("HDD", "Horizontal Directional Drilling – A construction method"),
    ("UG", "Underground – Construction methodology"),
    ("SHQ", "State Headquarters"),
    ("NLD", "National Long Distance – Network type prefix"),
]

table = doc.add_table(rows=1 + len(glossary), cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = "Abbreviation"
table.rows[0].cells[1].text = "Description"
for p in table.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True
for i, (abbr, desc) in enumerate(glossary):
    table.rows[i + 1].cells[0].text = abbr
    table.rows[i + 1].cells[1].text = desc

doc.add_paragraph("")

add_heading_styled(doc, "B. Key System URLs", level=2)

add_body_text(doc, "FTTX Home: https://ospreplica.jio.com/Runtime/Runtime/Form/FTTX__HOME/?")
add_body_text(doc, "Task List: https://ospreplica.jio.com/Runtime/Runtime/Form/FTTXTaskList/")

# ============================================================
# SAVE THE DOCUMENT
# ============================================================
doc.save(OUTPUT_DOC)
print(f"Word document created successfully: {OUTPUT_DOC}")
